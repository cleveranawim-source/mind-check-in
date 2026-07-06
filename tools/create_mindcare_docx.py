from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/교사용 마음건강 체크 도구 설명.docx"


COLORS = {
    "green": "547465",
    "deep": "25312C",
    "muted": "68736D",
    "fill": "F6F1E8",
    "line": "DCD3C3",
    "soft": "FBF8F1",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="DCD3C3", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    values = {"top": top, "start": start, "bottom": bottom, "end": end}
    for key, value in values.items():
        node = tc_mar.find(qn("w:" + key))
        if node is None:
            node = OxmlElement("w:" + key)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def paragraph_style(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    paragraph_style(p, before=14 if level == 1 else 9, after=6, line=1.1)
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, bold=True, color=COLORS["green"])
    elif level == 2:
        set_run_font(r, size=13, bold=True, color=COLORS["green"])
    else:
        set_run_font(r, size=12, bold=True, color=COLORS["deep"])
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    paragraph_style(p, before=0, after=6, line=1.12)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=COLORS["deep"])
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    paragraph_style(p, before=0, after=4, line=1.12)
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], size=10.5, color=COLORS["deep"])
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=COLORS["deep"])
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["fill"])
    set_cell_border(cell, color="D8CCB8", size="8")
    set_cell_margins(cell, top=160, bottom=160, start=200, end=200)
    p = cell.paragraphs[0]
    paragraph_style(p, before=0, after=3, line=1.1)
    r = p.add_run(label)
    set_run_font(r, size=10, bold=True, color=COLORS["green"])
    p2 = cell.add_paragraph()
    paragraph_style(p2, before=0, after=0, line=1.16)
    r2 = p2.add_run(text)
    set_run_font(r2, size=11, color=COLORS["deep"])
    spacer = doc.add_paragraph()
    paragraph_style(spacer, before=0, after=6)


def add_type_card(doc, title, summary, guidance):
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [2300, 7060])
    left, right = table.rows[0].cells
    set_cell_shading(left, "EEF3EF")
    set_cell_shading(right, "FFFFFF")
    for cell in (left, right):
        set_cell_border(cell, color="DCD3C3", size="6")
        set_cell_margins(cell, top=130, bottom=130, start=160, end=160)
    p = left.paragraphs[0]
    paragraph_style(p, before=0, after=0, line=1.1)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=COLORS["green"])
    p2 = right.paragraphs[0]
    paragraph_style(p2, before=0, after=4, line=1.12)
    r2 = p2.add_run(summary)
    set_run_font(r2, size=10.5, bold=True, color=COLORS["deep"])
    p3 = right.add_paragraph()
    paragraph_style(p3, before=0, after=0, line=1.12)
    r3 = p3.add_run(guidance)
    set_run_font(r3, size=10.2, color=COLORS["muted"])
    spacer = doc.add_paragraph()
    paragraph_style(spacer, before=0, after=4)


def add_type_detail(doc, title, understand, helpful_words, today_actions, week_actions):
    add_heading(doc, title, 2)
    add_body(doc, understand)
    add_callout(doc, "도움이 되는 말", helpful_words)
    add_heading(doc, "오늘 해볼 행동", 3)
    for item in today_actions:
        add_bullet(doc, item)
    add_heading(doc, "이번 주 조정", 3)
    for item in week_actions:
        add_bullet(doc, item)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Malgun Gothic"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
normal.font.size = Pt(10.5)

for name in ("List Bullet", "List Number"):
    style = styles[name]
    style.font.name = "Malgun Gothic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    style.font.size = Pt(10.5)

title = doc.add_paragraph()
paragraph_style(title, before=0, after=3, line=1.08)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("교사용 마음건강 체크 도구 설명")
set_run_font(r, size=24, bold=True, color=COLORS["deep"])

subtitle = doc.add_paragraph()
paragraph_style(subtitle, before=0, after=14, line=1.15)
r = subtitle.add_run("선생님들께 안내할 수 있는 검사 구조, 신뢰성 근거, 6가지 유형 해설")
set_run_font(r, size=12, color=COLORS["muted"])

add_callout(
    doc,
    "핵심 안내 문장",
    "이 검사는 선생님을 평가하거나 진단하기 위한 것이 아닙니다. 최근 2주 동안 내 마음이 어떤 방식으로 지쳐 있는지 확인하고, 지금 나에게 필요한 회복 방향을 찾기 위한 도구입니다. 결과는 낙인이 아니라 대화의 출발점입니다.",
)

add_heading(doc, "1. 이 도구는 무엇인가", 1)
add_body(
    doc,
    "이 도구는 교사의 마음건강을 진단하기 위한 검사가 아니라, 최근 2주 동안 내 마음과 몸이 어떤 방향으로 지쳐 있는지 살펴보는 자가 점검 도구입니다.",
)
add_body(
    doc,
    "학교 현장에서 교사는 단순히 스트레스를 받는다는 말로 설명하기 어려운 다양한 부담을 경험합니다. 수업 준비, 생활지도, 민원, 관계 갈등, 조직적 보호감의 부족, 교육적 신념의 흔들림, 특정 사건 이후의 긴장 등이 함께 작용합니다.",
)
add_body(
    doc,
    "따라서 이 도구는 선생님의 상태를 하나의 점수로만 보지 않고, 여섯 가지 마음건강 신호로 나누어 살펴봅니다. 결과는 '당신은 이런 사람입니다'가 아니라 '지금은 이런 회복 과제가 두드러집니다'라는 의미로 이해합니다.",
)

add_heading(doc, "2. 검사 구조", 1)
add_body(doc, "검사는 총 36문항으로 구성되어 있으며, 각 문항은 최근 2주를 기준으로 답합니다. 응답은 전혀 아님, 가끔, 보통, 자주, 거의 매일의 5단계입니다.")
for item in [
    "여섯 영역: 소진/에너지 고갈, 불안/과각성, 우울/위축, 분노/예민함, 무력감/도덕적 상처, 침습 기억/위협 반응",
    "문항 옆에는 영역명을 표시하지 않아 응답자가 특정 결과를 의식하고 답하는 편향을 줄입니다.",
    "가장 높은 영역은 주 유형, 두 번째로 높은 영역은 보조 신호로 제시합니다.",
    "위기 신호는 총점과 별도로 확인하며, 자해·자살 생각이나 타해 우려가 있으면 즉시 도움 연결을 우선합니다.",
]:
    add_bullet(doc, item)

add_heading(doc, "3. 신뢰성을 높이기 위한 설계", 1)
add_body(
    doc,
    "이 도구는 아직 임상적으로 표준화된 검사가 아니므로 진단명, 치료 여부, 병가 여부를 결정하는 근거로 단독 사용해서는 안 됩니다. 다만 아래 원칙을 반영해 선별 도구로서의 신뢰성을 높였습니다.",
)
for item in [
    "공공 정신건강 자가검진처럼 위험 영역을 먼저 선별하고, 필요 시 더 정밀한 평가로 이어지는 구조를 따릅니다.",
    "소진 영역은 WHO ICD-11에서 설명하는 직업적 소진 개념, 즉 에너지 고갈, 일과의 심리적 거리감, 직무 효능감 저하를 참고했습니다.",
    "교사 직무의 특수성을 반영해 민원 불안, 관계 긴장, 교육적 신념 손상, 사건 후유 반응을 별도 축으로 분리했습니다.",
    "위기 신호는 전체 점수와 상관없이 별도로 우선 확인합니다.",
]:
    add_bullet(doc, item)

doc.add_page_break()
add_heading(doc, "4. 여섯 가지 유형", 1)
types = [
    {
        "title": "회복 고갈형",
        "understand": "좋은 교사로 버티느라 몸과 마음의 연료가 먼저 소진된 상태입니다. 쉬어도 회복되지 않고 수업 전부터 기운이 빠지며, 학생에게 따뜻하게 반응할 여유가 줄어들 수 있습니다. 이 상태는 의지 부족이 아니라 회복 자원이 고갈되었다는 신호로 이해하는 것이 중요합니다.",
        "helpful_words": "지금 필요한 것은 더 세게 버티는 것이 아니라, 회복할 수 있는 조건을 다시 만드는 것입니다. 잘하고 싶은 마음이 컸기 때문에 여기까지 온 것일 수 있습니다.",
        "today": [
            "오늘 반드시 해야 할 일을 3개만 남기고 나머지는 내일 목록으로 옮깁니다.",
            "퇴근 후 학교 메신저나 업무 확인 시간을 한 번으로 제한합니다.",
            "수업 준비의 기준을 '완벽함'이 아니라 '운영 가능한 정도'로 낮춥니다.",
        ],
        "week": [
            "반복 업무 중 하나를 동료와 나누거나 관리자에게 조정 요청합니다.",
            "식사, 수면, 움직임 중 가장 무너진 한 가지를 먼저 회복 목표로 잡습니다.",
            "2주 이상 회복되지 않거나 결근 충동이 강하면 상담 또는 진료 가능성을 검토합니다.",
        ],
    },
    {
        "title": "경계 과부하형",
        "understand": "문제가 생기기 전에 막으려다 긴장 시스템이 계속 켜진 상태입니다. 민원, 연락, 관리자 호출, 실수 가능성에 민감해지고, 메시지를 여러 번 고치거나 퇴근 후에도 걱정이 이어질 수 있습니다. 불안은 위험을 예측하려는 보호 기능이지만, 너무 오래 켜져 있으면 삶 전체를 지치게 합니다.",
        "helpful_words": "불안하다는 것은 내가 약하다는 뜻이 아니라, 책임져야 할 것이 너무 많아졌다는 뜻일 수 있습니다. 모든 가능성을 혼자 예측하려 하지 않아도 됩니다.",
        "today": [
            "걱정되는 일을 '사실', '추정', '지금 할 수 있는 대응'으로 나누어 적습니다.",
            "보호자나 학생에게 보낼 문장은 초안 작성 후 10분 뒤 한 번만 다시 확인합니다.",
            "불안이 올라올 때 숨을 길게 내쉬며 어깨와 턱의 힘을 90초 동안 풀어봅니다.",
        ],
        "week": [
            "반복 민원이나 민감한 사안에 대한 표준 답변 문구를 미리 만들어 둡니다.",
            "혼자 판단하기 어려운 사안은 기록-공유-확인 순서로 처리합니다.",
            "불안을 줄이기 위해 반복하던 확인 행동 하나를 정해 조금씩 줄입니다.",
        ],
    },
    {
        "title": "의미 저하형",
        "understand": "교직의 보람, 흥미, 자기효능감이 낮아진 상태입니다. 학생의 긍정적인 반응도 마음에 잘 들어오지 않고, '내가 부족해서 그렇다'는 자기비난이 커질 수 있습니다. 단순한 피로보다 우울·위축 신호에 가까울 수 있으므로 오래 지속될 때는 가볍게 넘기지 않는 것이 좋습니다.",
        "helpful_words": "지금 보람이 느껴지지 않는다고 해서 선생님이 교사로서 실패한 것은 아닙니다. 마음이 너무 지쳐서 좋은 장면을 받아들이는 힘이 약해진 것일 수 있습니다.",
        "today": [
            "자기비난 문장을 사실 문장으로 바꿔 적습니다. 예: '나는 실패했다'가 아니라 '요즘 수업 후 회복이 어렵다'.",
            "하루 중 반드시 끝낼 수 있는 작은 행동 하나를 정하고 완료합니다.",
            "혼자 견디지 않기 위해 신뢰할 수 있는 사람 한 명에게 현재 상태를 짧게 알립니다.",
        ],
        "week": [
            "수업에서 작게라도 괜찮았던 장면을 하루 한 줄만 기록합니다.",
            "수면, 식사, 움직임을 간단히 기록해 생체 리듬이 얼마나 흔들렸는지 확인합니다.",
            "우울감, 흥미 저하, 무가치감이 2주 이상 이어지면 전문 평가를 권합니다.",
        ],
    },
    {
        "title": "감정 압력형",
        "understand": "참아온 억울함과 긴장이 분노나 예민함으로 새어 나오는 상태입니다. 말투가 날카로워지고 퇴근 후에도 특정 학생, 보호자, 사건에 대한 감정이 가라앉지 않을 수 있습니다. 분노는 없애야 할 감정이라기보다 침해된 경계와 반복된 무력감의 신호일 때가 많습니다.",
        "helpful_words": "화가 난다는 사실만으로 나쁜 교사가 되는 것은 아닙니다. 중요한 것은 감정을 부정하는 것이 아니라, 관계를 다치게 하기 전에 안전하게 표현하는 방법을 찾는 것입니다.",
        "today": [
            "즉시 답하지 않아도 되는 상황이면 20분 뒤에 반응하기로 정합니다.",
            "감정을 '화남' 하나로 묶지 말고 억울함, 무시당함, 두려움, 피로처럼 구체적으로 이름 붙입니다.",
            "지금 해결할 문제와 나중에 다룰 문제를 분리합니다.",
        ],
        "week": [
            "반복해서 분노가 올라오는 상황의 공통 조건을 찾아봅니다.",
            "학생·보호자 대응에서 사용할 짧고 단호한 문장을 미리 준비합니다.",
            "사과가 필요한 상황과 경계 설정이 필요한 상황을 구분해 동료와 함께 점검합니다.",
        ],
    },
    {
        "title": "도덕적 상처형",
        "understand": "교육적 신념과 학교 현실이 충돌하면서 마음이 다친 상태입니다. 옳다고 믿는 교육적 판단을 지키기 어렵거나, 학교가 나를 보호하지 못한다는 느낌이 강할 수 있습니다. 이 유형은 개인의 마음관리만으로 해결하기 어렵고, 절차와 보호, 동료 지지, 조직적 대응이 함께 필요합니다.",
        "helpful_words": "선생님이 약해서 흔들리는 것이 아니라, 중요하게 여기는 가치가 반복해서 손상되었기 때문에 아픈 것일 수 있습니다. 이 고통은 개인의 성격 문제가 아니라 학교 환경과도 연결되어 있습니다.",
        "today": [
            "내 책임과 시스템 책임을 분리해 적습니다.",
            "사건이나 상황을 감정 평가 없이 사실 중심으로 기록합니다.",
            "혼자 판단하지 말고 신뢰할 수 있는 동료나 공식 라인에 공유합니다.",
        ],
        "week": [
            "교권보호, 민원 대응, 상담 지원 등 활용 가능한 절차를 확인합니다.",
            "지금도 지킬 수 있는 나의 교육적 기준을 아주 작은 단위로 다시 정합니다.",
            "조직적 보호가 필요한 사안은 개인 감정 문제가 아니라 절차 문제로 다룹니다.",
        ],
    },
    {
        "title": "사건 후유 반응형",
        "understand": "위협적이거나 충격적인 사건 이후 몸과 기억이 아직 안전을 확인하지 못한 상태입니다. 특정 장면이나 말이 갑자기 떠오르고, 비슷한 연락이나 공간을 피하고 싶어지며, 수면과 집중이 흔들릴 수 있습니다. 머리로는 끝난 일이라고 알아도 몸은 아직 경계할 수 있습니다.",
        "helpful_words": "아직도 반응한다고 해서 유난한 것이 아닙니다. 몸과 마음이 위험했던 순간을 기억하고 다시 안전해지려는 과정일 수 있습니다.",
        "today": [
            "지금 있는 장소, 날짜, 발바닥 감각을 말하며 현재로 돌아오는 연습을 합니다.",
            "사건 관련 연락이나 공간 노출을 혼자 감당하지 않도록 동행이나 조정을 요청합니다.",
            "수면과 식사를 회복 목표의 1순위로 둡니다.",
        ],
        "week": [
            "사건 기록과 감정 기록을 분리해 정리합니다.",
            "학교 내 안전 동선과 대응 동반자를 정합니다.",
            "악몽, 플래시백, 회피, 과각성이 이어지면 트라우마 평가나 상담을 권합니다.",
        ],
    },
]
for idx, item in enumerate(types):
    if idx in (2, 4):
        doc.add_page_break()
    add_type_detail(doc, item["title"], item["understand"], item["helpful_words"], item["today"], item["week"])

doc.add_page_break()
add_heading(doc, "5. 선생님들께 설명할 때의 흐름", 1)
for item in [
    "먼저 이 도구가 진단이나 평가가 아니라 자기이해를 위한 점검임을 분명히 안내합니다.",
    "최근 2주를 기준으로 답하도록 요청하고, 특정 결과를 의식하지 않고 편하게 답하도록 설명합니다.",
    "결과는 주 유형과 보조 신호로 읽고, 한 사람을 하나의 유형으로 고정하지 않는다고 안내합니다.",
    "위기 신호가 있으면 점수 해석보다 즉시 사람과 전문기관에 연결하는 것이 우선이라고 말합니다.",
]:
    add_bullet(doc, item)

add_heading(doc, "6. 참고 근거", 1)
add_body(doc, "WHO ICD-11 Burn-out: https://www.who.int/standards/classifications/frequently-asked-questions/burn-out-an-occupational-phenomenon")
add_body(doc, "국가정신건강정보포털 자가검진: https://www.mentalhealth.go.kr/portal/mdexmnDtl/mdexmnTypeList.do")
add_body(doc, "보건복지부 109 자살예방상담전화 안내: https://www.mohw.go.kr/board.es?act=view&bid=0027&list_no=1479607&mid=a10503010100&nPage=1&tag=")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph_style(footer, before=0, after=0, line=1.0)
fr = footer.add_run("LevLab")
set_run_font(fr, size=9, bold=True, color="8C9992")

doc.save(OUT)
print(OUT)
